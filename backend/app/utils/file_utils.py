import io
import os
import aiofiles
from typing import Optional
from fastapi import UploadFile
from PyPDF2 import PdfReader, errors as PyPDF2Errors
import docx

# -----------------------------
# Save uploaded file async
# -----------------------------
async def save_upload_file(upload_file: UploadFile, destination: str) -> None:
    # IMPORTANT: Reset cursor to start before reading
    await upload_file.seek(0)
    async with aiofiles.open(destination, "wb") as out_file:
        while content := await upload_file.read(1024):
            await out_file.write(content)
    # Reset cursor again so other functions can read it later
    await upload_file.seek(0)

# -----------------------------
# Extraction Logic (Kept exactly as requested)
# -----------------------------

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() or ""
        
        print(f"DEBUG: Extracted {len(text)} characters from {file_path}")
        
        if len(text.strip()) == 0:
            print("WARNING: Extraction returned NO text. Is this a scanned image?") 
            
    except Exception as e:
        print(f"PDF text extraction error: {e}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        print(f"DEBUG: Extracted {len(text)} characters from {file_path}")
        
        if len(text.strip()) == 0:
            print("WARNING: Extraction returned NO text. Is this a scanned image?") 
            
    except Exception as e:
        print(f"DOCX text extraction error: {e}")
    return text

def extract_text(file_path: str, filename: str) -> Optional[str]:
    filename_low = filename.lower()
    if filename_low.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif filename_low.endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif filename_low.endswith(".txt"):
        try:
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            print(f"TXT file read error: {e}")
            return None
    return None

def extract_text_from_memory(file_content: bytes, filename: str) -> Optional[str]:
    """
    Extracts text from byte content (Your exact logic).
    """
    text = ""
    filename_low = filename.lower()
    try:
        if filename_low.endswith(".pdf"):
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                text += page.extract_text() or ""
        
        elif filename_low.endswith(".docx"):
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        elif filename_low.endswith(".txt"):
            text = file_content.decode("utf-8", errors="ignore")
        
        else:
            return None

        result = text.strip()
        if not result:
            print(f"DEBUG: Extracted text is empty for {filename}")
            return None
        return result

    except PyPDF2Errors.PdfReadError as e:
        print(f"ERROR: Could not read PDF file '{filename}': {e}")
        return None
    except Exception as e:
        print(f"ERROR: Unexpected error during extraction for '{filename}': {e}")
        return None